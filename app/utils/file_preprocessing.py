import pandas as pd
import json
from typing import Union, BinaryIO, Tuple
from pathlib import Path
import docx
import requests
from PIL import Image
import io
from app.utils.file_management import temp_file_manager
import fitz  # PyMuPDF
from app.utils.vision_processing import VisionProcessor
from tempfile import SpooledTemporaryFile
from typing import Dict, Tuple
from app.utils.data_processing import sanitize_error_message, get_data_snapshot
from fastapi import UploadFile
from app.schemas import FileDataInfo, FileMetadata
from typing import List
import logging
from app.utils.google_integration import GoogleIntegration
from app.utils.microsoft_integration import MicrosoftIntegration
from supabase import create_client, Client as SupabaseClient
import asyncio
import aiohttp


class FilePreprocessor:
    """Handles preprocessing of various file types for data processing pipeline."""
    
    def __init__(self, num_images_processed: int = 0, supabase: SupabaseClient = None, user_id: str = None):
        """Initialize the FilePreprocessor with optional authentication"""
        self.num_images_processed = num_images_processed
        self.supabase = supabase
        self.user_id = user_id
        self.google_integration = None
        self.microsoft_integration = None
        
        if supabase and user_id:
            self.google_integration = GoogleIntegration(supabase, user_id)
            self.microsoft_integration = MicrosoftIntegration(supabase, user_id)

    @staticmethod
    def process_excel(file: Union[SpooledTemporaryFile, str, Path, BinaryIO]) -> pd.DataFrame:
        """
        Process Excel files (.xlsx) and convert to pandas DataFrame
        
        Args:
            file: File object (SpooledTemporaryFile, BytesIO) or path to Excel file
            
        Returns:
            pd.DataFrame: Processed data as DataFrame
        """
        try:
            print("Attempting to process excel file")
            # For file paths
            if isinstance(file, (str, Path)):
                return pd.read_excel(file)
            
            # For file objects (BytesIO, SpooledTemporaryFile)
            if hasattr(file, 'seek'):
                file.seek(0)
                # Create a BytesIO object from the file content
                content = file.read()
                if isinstance(content, str):
                    raise ValueError("File content must be bytes")
                return pd.read_excel(io.BytesIO(content))
            
            raise ValueError("Unsupported file object type")
        except Exception as e:
            raise ValueError(f"Error processing Excel file: {str(e)}")

    @staticmethod
    def process_csv(file: Union[BinaryIO, str]) -> pd.DataFrame:
        """
        Process CSV files and convert to pandas DataFrame
        
        Args:
            file: File object or path to CSV file
            
        Returns:
            pd.DataFrame: Processed data as DataFrame
        """
        try:
            return pd.read_csv(file)
        except Exception as e:
            raise ValueError(FilePreprocessor._sanitize_error(e))

    @staticmethod
    def process_json(file: Union[BinaryIO, str]) -> str:
        """
        Process JSON files and convert to string
        
        Args:
            file: File object or path to JSON file
            
        Returns:
            str: JSON content as string
        """
        try:
            if isinstance(file, str):
                with open(file, 'r') as f:
                    data = json.load(f)
            else:
                data = json.load(file)
            return json.dumps(data)
        except Exception as e:
            raise ValueError(FilePreprocessor._sanitize_error(e))

    @staticmethod
    def process_text(file: Union[BinaryIO, str]) -> str:
        """
        Process text files (.txt) and convert to string
        
        Args:
            file: File object or path to text file
            
        Returns:
            str: File content as string
        """
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        try:
            if isinstance(file, str):
                for encoding in encodings:
                    try:
                        with open(file, 'r', encoding=encoding) as f:
                            return f.read()
                    except UnicodeDecodeError:
                        continue
            else:
                content = file.read()
                for encoding in encodings:
                    try:
                        return content.decode(encoding)
                    except UnicodeDecodeError:
                        continue
                
                # If all encodings fail, try with error handling
                return content.decode('utf-8', errors='replace')
                
            raise ValueError("Unable to decode file with any supported encoding")
        except Exception as e:
            raise ValueError(FilePreprocessor._sanitize_error(e))

    @staticmethod
    def process_docx(file: Union[BinaryIO, str]) -> str:
        """
        Process Word documents (.docx) and convert to string
        
        Args:
            file: File object or path to Word document
            
        Returns:
            str: Document content as string
        """
        try:
            if isinstance(file, str):
                doc = docx.Document(file)
            else:
                doc = docx.Document(io.BytesIO(file.read()))
            return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        except Exception as e:
            raise ValueError(FilePreprocessor._sanitize_error(e))

    def process_image(self, file: Union[BinaryIO, str], output_path: str = None, query: str = None) -> Tuple[str, str]:
        """
        Process image files (.png, .jpg, .jpeg) and extract content using vision processing
        
        Args:
            file: File object or path to file
            output_path: Optional path to save converted image
            query: Optional query for vision processing
            
        Returns:
            Tuple[str, str]: (vision_content, new_file_path)
        """
        try:
            # Process the image file
            if isinstance(file, str):
                img = Image.open(file)
                image_path = file
            else:
                try:
                    if hasattr(file, 'read'):
                        content = file.read()
                        if not isinstance(content, bytes):
                            raise ValueError("Invalid image content")
                        img = Image.open(io.BytesIO(content))
                    else:
                        img = Image.open(file)
                except Exception:
                    raise ValueError("Unable to read image file")

            # Convert PNG to JPEG if needed
            new_path = None
            if img.format == 'PNG' and output_path:
                if img.mode in ('RGBA', 'P'):
                    img = img.convert('RGB')
                img.save(output_path, 'JPEG')
                new_path = output_path
                image_path = new_path
            else:
                # Save the original file if it wasn't converted
                image_path = output_path or str(Path(output_path).parent / "original_image")
                with open(image_path, 'wb') as f:
                    if isinstance(file, str):
                        with open(file, 'rb') as src:
                            f.write(src.read())
                    else:
                        file.seek(0)
                        f.write(file.read())

            # Process with vision
            vision_processor = VisionProcessor()
            vision_result = vision_processor.process_image_with_vision(
                image_path=image_path,
                query=query
            )

            if vision_result["status"] == "error":
                raise ValueError(f"Vision API error: {vision_result['error']}")

            # Increment the image counter
            self.num_images_processed += 1

            return vision_result["content"]

        except Exception as e:
            raise ValueError(FilePreprocessor._sanitize_error(e))

    @staticmethod
    def process_gsheet_url(url: str, supabase: SupabaseClient = None, user_id: str = None) -> pd.DataFrame:
        """
        Process Google Sheets URLs and convert to pandas DataFrame
        
        Args:
            url (str): URL of the Google Sheet file
            supabase (SupabaseClient): Supabase client for authentication
            user_id (str): User ID for authentication
            
        Returns:
            pd.DataFrame: DataFrame containing the sheet data
            
        Raises:
            ValueError: If URL is invalid or file cannot be accessed
        """
        try:
            if not supabase or not user_id:
                raise ValueError("Authentication required to access Google Sheets")

            # Initialize Google integration
            g_integration = GoogleIntegration(supabase, user_id)
            
            if 'docs.google.com/spreadsheets' in url:
                try:
                    # Extract file ID from URL
                    file_id = url.split('/d/')[1].split('/')[0]
                    
                    # Use Google integration to get authenticated service
                    service = g_integration.get_sheets_service()
                    if not service:
                        raise ValueError("Failed to initialize Google Sheets service")
                    
                    # Read the sheet data using authenticated service
                    result = service.spreadsheets().values().get(
                        spreadsheetId=file_id,
                        range='A:ZZ'  # Get all data
                    ).execute()
                    
                    # Convert to DataFrame
                    values = result.get('values', [])
                    if not values:
                        return pd.DataFrame()
                    
                    headers = values[0]
                    data = values[1:]
                    return pd.DataFrame(data, columns=headers)
                    
                except Exception as e:
                    raise ValueError(f"Failed to access Google Sheet: {str(e)}")
            else:
                raise ValueError("Invalid Google Sheets URL format")
                
        except Exception as e:
            error_msg = str(e)
            if any(sensitive in error_msg.lower() for sensitive in ['token', 'key', 'auth', 'password', 'secret']):
                error_msg = "Authentication error occurred while accessing the file"
            raise ValueError(FilePreprocessor._sanitize_error(error_msg))

    @staticmethod
    async def process_msft_excel_url(url: str, supabase: SupabaseClient = None, user_id: str = None) -> pd.DataFrame:
        """
        Process Microsoft Excel Online URLs and convert to pandas DataFrame
        
        Args:
            url (str): URL of the Microsoft Excel Online file
            supabase (SupabaseClient): Supabase client for authentication
            user_id (str): User ID for authentication
            
        Returns:
            pd.DataFrame: DataFrame containing the sheet data
            
        Raises:
            ValueError: If URL is invalid or file cannot be accessed
        """
        try:
            if not supabase or not user_id:
                raise ValueError("Authentication required to access Microsoft Excel files")

            # Initialize Microsoft integration
            msft_integration = MicrosoftIntegration(supabase, user_id)
            
            if any(domain in url.lower() for domain in ['onedrive.live.com', 'sharepoint.com']):
                try:
                    # Get drive ID and item ID from URL
                    drive_id, item_id = await msft_integration._get_one_drive_and_item_info(url)
                    
                    # Get authenticated headers
                    headers = await msft_integration._get_microsoft_headers()
                    
                    # Create a session for the workbook
                    session_id = await msft_integration._manage_office_session(item_id, 'create')
                    
                    try:
                        # Get the first worksheet
                        async with aiohttp.ClientSession() as session:
                            async with session.get(
                                f'https://graph.microsoft.com/v1.0/drives/{drive_id}/items/{item_id}/workbook/worksheets',
                                headers=headers
                            ) as response:
                                if response.status != 200:
                                    raise ValueError("Failed to access Excel worksheet")
                                worksheets = await response.json()
                                if not worksheets['value']:
                                    raise ValueError("No worksheets found in the workbook")
                                
                                active_sheet = worksheets['value'][0]
                                
                                # Get the used range to read all data
                                async with session.get(
                                    f'https://graph.microsoft.com/v1.0/drives/{drive_id}/items/{item_id}/workbook/worksheets/{active_sheet["id"]}/usedRange',
                                    headers=headers
                                ) as range_response:
                                    if range_response.status != 200:
                                        raise ValueError("Failed to read worksheet data")
                                    
                                    range_data = await range_response.json()
                                    values = range_data.get('values', [])
                                    
                                    if not values:
                                        return pd.DataFrame()
                                    
                                    # Convert to DataFrame
                                    headers = values[0]
                                    data = values[1:]
                                    return pd.DataFrame(data, columns=headers)
                    
                    finally:
                        # Always close the session
                        if session_id:
                            await msft_integration._manage_office_session(session_id, 'close')
                            
                except Exception as e:
                    raise ValueError(f"Failed to process Excel Online file: {str(e)}")
            else:
                raise ValueError("Invalid Microsoft Excel Online URL format")
                
        except Exception as e:
            error_msg = str(e)
            if any(sensitive in error_msg.lower() for sensitive in ['token', 'key', 'auth', 'password', 'secret']):
                error_msg = "Authentication error occurred while accessing the file"
            raise ValueError(FilePreprocessor._sanitize_error(error_msg))

    def process_pdf(self, file: Union[BinaryIO, str], query: str = None) -> Tuple[str, str, bool]:
        """
        Process PDF files and convert to string if readable, otherwise handle with vision
        
        Args:
            file: File object or path to PDF file
            query: Optional query for vision processing
            
        Returns:
            Tuple[str, str, bool]: (content, data_type, is_readable)
            - content: Extracted text or vision API result
            - data_type: "text" or "vision_extracted"
            - is_readable: Whether the PDF was machine-readable
        """
        try:
            # Create a temporary file if we received a file object
            if not isinstance(file, str):
                temp_path = temp_file_manager.save_temp_file(file.read(), "temp.pdf")
                pdf_path = str(temp_path)
            else:
                pdf_path = file

            # Open PDF with PyMuPDF
            doc = fitz.open(pdf_path)
            
            # Try to extract text
            text_content = ""
            total_text_length = 0
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                page_text = page.get_text()
                total_text_length += len(page_text.strip())
                text_content += f"\n[Page {page_num + 1}]\n{page_text}"

            doc.close()

            # If not a file path, clean up the temporary file
            if not isinstance(file, str):
                temp_path.unlink()

            # Check if PDF is readable (has meaningful text content)
            is_readable = total_text_length > 100  # Arbitrary threshold

            if is_readable:
                return text_content, "text", True
            
            # Handle unreadable PDFs
            if doc.page_count > 5:
                raise ValueError("Unreadable PDF with more than 5 pages")
                
            # For small unreadable PDFs, process with vision
            vision_processor = VisionProcessor()
            vision_result = vision_processor.process_pdf_with_vision(pdf_path, query)
            
            if vision_result["status"] == "error":
                raise ValueError(f"Vision API error: {vision_result['error']}")

            # Increment the image counter for each page in unreadable PDF
            self.num_images_processed += doc.page_count

                
            return vision_result["content"], "vision_extracted", False

        except Exception as e:
            raise ValueError(FilePreprocessor._sanitize_error(e))

    def preprocess_file(self, file: Union[BinaryIO, str], file_type: str, supabase: SupabaseClient = None, user_id: str = None) -> Union[str, pd.DataFrame]:
        """
        Preprocess file based on its type
        """
        processors = {
            'xlsx': self.process_excel,
            'csv': self.process_csv,
            'json': self.process_json,
            'txt': self.process_text,
            'docx': self.process_docx,
            'png': self.process_image,
            'jpg': self.process_image,
            'jpeg': self.process_image,
            'pdf': self.process_pdf,
            'gsheet': lambda url: self.process_gsheet_url(url, self.supabase, self.user_id),
            'office_sheet': lambda url: asyncio.run(self.process_msft_excel_url(url, self.supabase, self.user_id))
        }
        
        processor = processors.get(file_type.lower())
        if not processor:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        return processor(file)

    @staticmethod
    def _sanitize_error(e: Exception) -> str:
        """Sanitize error messages to prevent binary data leakage"""
        try:
            error_msg = str(e)
            # If message contains binary data or is too long, return generic message
            if not error_msg.isascii() or len(error_msg) > 200:
                return f"Error processing file: {e.__class__.__name__}"
            return error_msg.encode('ascii', 'ignore').decode('ascii')
        except:
            return "Error processing file"


def preprocess_files(
    files: List[UploadFile],
    files_metadata: List[FileMetadata],
    web_urls: List[str],
    query: str,
    session_dir,
    supabase: SupabaseClient,
    user_id: str,
    num_images_processed: int = 0
) -> Tuple[List[FileDataInfo], int]:
    """Helper function to preprocess files and web URLs"""
    preprocessor = FilePreprocessor(num_images_processed, supabase, user_id)
    processed_data = []
    
    # Process web URLs if provided
    for url in web_urls:
        try:
            logging.info(f"Processing URL: {url}")
            if 'docs.google' in url:
                content = preprocessor.process_gsheet_url(url, 'gsheet', supabase, user_id) 
            elif 'onedrive.live' in url:
                content = preprocessor.process_msft_excel_url(url, 'office_sheet', supabase, user_id)
            else:
                logging.error(f"Unsupported URL format: {url}")
                continue
            
            data_type = "DataFrame" if isinstance(content, pd.DataFrame) else "text"
            processed_data.append(
                FileDataInfo(
                    content=content,
                    snapshot=get_data_snapshot(content, data_type),
                    data_type=data_type,
                    original_file_name=url.split('/')[-1],
                    url=url
                )
            )
        except Exception as e:
            error_msg = sanitize_error_message(e)
            logging.error(f"Error processing URL {url}: {error_msg}")
            raise Exception(f"Error processing URL {url}: {error_msg}")
    
    # Process uploaded files using metadata
    if files and files_metadata:
        for metadata in files_metadata:
            try:
                file = files[metadata.index]
                logging.info(f"Preprocessing file: {metadata.name} with type: {metadata.type}")
                
                # Map MIME types to FilePreprocessor types
                mime_to_processor = {
                    'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': 'xlsx',
                    'text/csv': 'csv',
                    'application/json': 'json',
                    'text/plain': 'txt',
                    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
                    'image/png': 'png',
                    'image/jpeg': 'jpg',
                    'image/jpg': 'jpg',
                    'application/pdf': 'pdf'
                }
                
                file_type = mime_to_processor.get(metadata.type)
                if not file_type:
                    raise ValueError(f"Unsupported MIME type: {metadata.type}")

                # Handle special cases for images and PDFs that need additional parameters
                kwargs = {}
                if file_type in ['png', 'jpg', 'jpeg']:
                    kwargs['output_path'] = str(session_dir / f"{metadata.name}.jpeg")
                elif file_type == 'pdf':
                    kwargs['query'] = query

                # Process the file
                with io.BytesIO(file.file.read()) as file_obj:
                    file_obj.seek(0)  # Reset the file pointer to the beginning
                    content = preprocessor.preprocess_file(file_obj, file_type, **kwargs)

                # Handle different return types
                if file_type == 'pdf':
                    content, data_type, is_readable = content  # Unpack PDF processor return values
                    metadata_info = {"is_readable": is_readable}
                else:
                    data_type = "DataFrame" if isinstance(content, pd.DataFrame) else "text"
                    metadata_info = {}

                processed_data.append(
                    FileDataInfo(
                        content=content,
                        snapshot=get_data_snapshot(content, data_type),
                        data_type=data_type,
                        original_file_name=metadata.name,
                        metadata=metadata_info,
                        new_file_path=kwargs.get('output_path')  # For images
                    )
                )

            except Exception as e:
                error_msg = sanitize_error_message(e)
                logging.error(f"Error processing file {metadata.name}: {error_msg}")
                raise ValueError(f"Error processing file {metadata.name}: {error_msg}")

    return processed_data, preprocessor.num_images_processed
